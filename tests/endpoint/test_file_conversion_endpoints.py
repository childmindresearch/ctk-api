"""Tests the diagnoses endpoints."""
import pathlib

import docx
import pytest
from fastapi import status, testclient

from . import conftest


@pytest.mark.asyncio()
async def test_post_md2docx(
    client: testclient.TestClient,
    endpoints: conftest.Endpoints,
    tmp_path: pathlib.Path,
) -> None:
    """Tests the anonymization endpoint."""
    md_text = "# Test Markdown Text"

    response = client.post(
        endpoints.POST_MARKDOWN_TO_DOCX,
        data={"markdown_text": md_text},
    )
    with (tmp_path / "md2docx.docx").open("wb") as output_file:
        output_file.write(response.content)
        document = docx.Document(tmp_path / "md2docx.docx")

    assert response.status_code == status.HTTP_200_OK
    assert (
        response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert (
        response.headers["Content-Disposition"] == 'attachment; filename="md2docx.docx"'
    )
    assert document.paragraphs[0].text == md_text[2:]
    assert document.paragraphs[0].style.name == "Heading 1"


@pytest.mark.asyncio()
async def test_post_intake2docx(
    client: testclient.TestClient,
    endpoints: conftest.Endpoints,
    tmp_path: pathlib.Path,
    data_dir: pathlib.Path,
) -> None:
    """Tests the intake to docx endpoint.

    This only tests whether the endpoint passes without error and checks for a single
    paragraph. The full functionality of the endpoint is tested in the unit tests.
    """
    intake_csv = (data_dir / "test_redcap_data.csv").open("rb")

    response = client.post(
        endpoints.POST_INTAKE_TO_DOCX,
        files={"csv_file": ("test_redcap_data.csv", intake_csv, "text/csv")},
        data={"redcap_survey_identifier": "1"},
    )
    with (tmp_path / "intake2docx.docx").open("wb") as output_file:
        output_file.write(response.content)
        document = docx.Document(tmp_path / "intake2docx.docx")

    assert response.status_code == status.HTTP_200_OK
    assert (
        response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert (
        response.headers["Content-Disposition"] == 'attachment; filename="report.docx"'
    )
    assert document.paragraphs[6].text == 'Name: Shizuka "Lea" Sakai'
