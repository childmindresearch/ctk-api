"""Configurations for all tests."""

import dataclasses
import pathlib
import tempfile
from collections.abc import Generator
from typing import Any

import docx
import pytest
from sqlalchemy import orm

from ctk_api.microservices import sql
from ctk_api.routers.file_conversion.intake import parser


@pytest.fixture(autouse=True, scope="session")
def _reset_testing_db() -> None:
    """Resets the testing database."""
    database = sql.Database()
    sql.Base.metadata.drop_all(database.engine)
    database.create_database_schema()


@pytest.fixture(scope="session")
def data_dir() -> pathlib.Path:
    """Gets the data directory."""
    return pathlib.Path(__file__).parent / "data"


@pytest.fixture()
def session() -> orm.Session:
    """Gets a database session."""
    return next(sql.get_session())


@pytest.fixture()
def document() -> Generator[tempfile._TemporaryFileWrapper, None, None]:
    """Generates a fixture Word document.

    This document contains a title and a paragraph.

    Returns:
        Generator[tempfile._TemporaryFileWrapper, None, None]: The saved file.
    """
    doc = docx.Document()
    doc.add_heading("Title", 0)
    doc.add_heading("clinical summary and impressions", 1)
    doc.add_paragraph("Name: Lea Avatar")
    doc.add_paragraph("He she herself man")
    doc.add_heading("recommendations", 1)
    doc.add_paragraph("hello!")

    with tempfile.NamedTemporaryFile(suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        yield temp_file


@pytest.fixture(scope="session")
def test_redcap_data(data_dir: pathlib.Path) -> dict[str, Any]:
    """Returns a dictionary of test data."""

    @dataclasses.dataclass
    class FastApiUploadFileMimic:
        file = data_dir / "test_redcap_data.csv"

    data_frame = parser.read_subject_row(FastApiUploadFileMimic, 1)  # type: ignore[arg-type]
    return data_frame.row(0, named=True)
