"""Tests for the intake writer module."""

import pathlib

import docx
import pytest

from ctk_api.routers.file_conversion.intake.utils import (
    docx_utils,
    language_utils,
    string_utils,
)


def test_find_replace_replaces_text(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text in a Word document."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = docx_utils.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].text.endswith("heading_replace")
    assert "paragraph_replace" in replace.document.paragraphs[1].text
    assert (
        replace.document.sections[0]
        .footer.paragraphs[0]
        .text.endswith("footer_replace")
    )
    assert (
        replace.document.sections[0]
        .header.paragraphs[0]
        .text.endswith("header_replace")
    )


def test_find_replace_maintains_style(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text maintains the style."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = docx_utils.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].style.name == "Heading 1"
    assert replace.document.paragraphs[1].style.name == "Normal"
    assert replace.document.sections[0].footer.paragraphs[0].style.name == "Footer"
    assert replace.document.sections[0].header.paragraphs[0].style.name == "Header"


def test_find_replace_maintains_custom_styling(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text maintains custom styling (e.g. bold).

    Note: all lines in the test file start with a bolded word.
    """
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = docx_utils.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].runs[0].bold
    assert replace.document.paragraphs[1].runs[0].bold
    assert replace.document.sections[0].footer.paragraphs[0].runs[0].bold
    assert replace.document.sections[0].header.paragraphs[0].runs[0].bold


def test_find_replace_overflow(
    data_dir: pathlib.Path,
) -> None:
    """Tests that find and replace does not remove additional text."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = docx_utils.DocxReplace(document)

    replace.replace("This is a paragraph", "paragraph_replace")

    assert (
        replace.document.paragraphs[1].text
        == "Paragraph: paragraph_replace with overflow."
    )


@pytest.mark.parametrize(
    ("rgb", "hexadecimal"),
    [
        ((0, 0, 0), "#000000"),
        ((255, 255, 255), "#FFFFFF"),
        ((255, 0, 0), "#FF0000"),
        ((0, 255, 0), "#00FF00"),
        ((0, 0, 255), "#0000FF"),
        ((255, 255, 0), "#FFFF00"),
        ((0, 255, 255), "#00FFFF"),
        ((255, 0, 255), "#FF00FF"),
        ((128, 128, 128), "#808080"),
        ((255, 128, 128), "#FF8080"),
        ((128, 255, 128), "#80FF80"),
        ((128, 128, 255), "#8080FF"),
        ((255, 255, 128), "#FFFF80"),
        ((128, 255, 255), "#80FFFF"),
        ((255, 128, 255), "#FF80FF"),
    ],
)
def test_rgb_to_hex(
    rgb: tuple[int, int, int],
    hexadecimal: str,
) -> None:
    """Tests converting RGB to hex."""
    assert docx_utils.rgb_to_hex(*rgb) == hexadecimal


@pytest.mark.parametrize(
    ("elements", "expected"),
    [
        ([], ""),
        (["a"], "a"),
        (["a", "b"], "a and b"),
        (["a", "b", "c"], "a, b, and c"),
        (["a", "b", "c", "d"], "a, b, c, and d"),
    ],
)
def test_join_with_oxford_comma(elements: list[str], expected: str) -> None:
    """Tests joining a list with an Oxford comma."""
    assert string_utils.join_with_oxford_comma(elements) == expected


@pytest.mark.parametrize(
    ("rank", "suffix"),
    [
        (-1, "th"),
        (0, "th"),
        (1, "st"),
        (2, "nd"),
        (3, "rd"),
        (4, "th"),
        (11, "th"),
        (12, "th"),
        (13, "th"),
        (14, "th"),
        (21, "st"),
        (22, "nd"),
        (23, "rd"),
        (24, "th"),
        (1111, "th"),
        ("1", "st"),
    ],
)
def test_ordinal_suffix(rank: int, suffix: str) -> None:
    """Tests getting the ordinal suffix of a number."""
    assert string_utils.ordinal_suffix(rank) == suffix


@pytest.mark.parametrize(
    ("sentence", "expected"),
    [
        ("They hears it in your voice.", "They hear it in your voice."),
        ('She hears it in your voice."', "She hears it in your voice."),
        (
            "They has been waiting from sprinkler splashes until fireplace ashes.",
            "They have been waiting from sprinkler splashes until fireplace ashes.",
        ),
        (
            "She has been waiting from sprinkler splashes until fireplace ashes.",
            "She has been waiting from sprinkler splashes until fireplace ashes.",
        ),
    ],
)
def correct_verb_conjugation(sentence: str, expected: str) -> None:
    """Tests whether the verb conjugations of they are corrected."""
    obj = language_utils.DocumentCorrections(document="")
    pairs = obj._find_subject_verb(sentence)

    for pair in pairs:
        sentence = obj._correct_they_verb_conjugation(
            sentence,
            pair,
        )

    assert sentence == expected


def test_remove_excess_whitespace() -> None:
    """Test the _remove_excess_whitespace method."""
    test_text = """    This   is a test string with

    excess    whitespace.     """
    expected = "This is a test string with excess whitespace."

    actual = string_utils.remove_excess_whitespace(test_text)

    assert actual == expected


@pytest.mark.parametrize(
    ("string", "expected"),
    [
        ("1", 1),
        ("1st", 1),
        ("1 month", 1),
        ("1.5", 1.5),
        ("1.5 months", 1.5),
        ("twenty five", 25),
        ("5 months and 4 days", "5 months and 4 days"),
    ],
)
def test_string_to_int(string: str, expected: float) -> None:
    """Test the string_to_int method."""
    actual = string_utils.StringToInt().parse(string)

    assert actual == expected
    assert isinstance(actual, type(expected))
