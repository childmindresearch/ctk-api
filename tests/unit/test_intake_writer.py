"""Tests for the intake writer module."""
import pathlib

import docx

from ctk_api.routers.file_conversion.intake import writer


def test_find_replace_replaces_text(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text in a Word document."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = writer.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].text.endswith("heading_replace")
    assert "paragraph_replace" in replace.document.paragraphs[1].text
    assert (
        replace.document.sections[0]
        .footer.paragraphs[0]
        .text.endswith("footer_replace")
    )
    assert (
        replace.document.sections[0]
        .header.paragraphs[0]
        .text.endswith("header_replace")
    )


def test_find_replace_maintains_style(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text maintains the style."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = writer.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].style.name == "Heading 1"
    assert replace.document.paragraphs[1].style.name == "Normal"
    assert replace.document.sections[0].footer.paragraphs[0].style.name == "Footer"
    assert replace.document.sections[0].header.paragraphs[0].style.name == "Header"


def test_find_replace_maintains_custom_styling(
    data_dir: pathlib.Path,
) -> None:
    """Tests finding and replacing text maintains custom styling (e.g. bold).

    Note: all lines in the test file start with a bolded word.
    """
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = writer.DocxReplace(document)

    replace.replace("This is a header", "header_replace")
    replace.replace("This is a footer", "footer_replace")
    replace.replace("This is a paragraph", "paragraph_replace")
    replace.replace("This is a heading", "heading_replace")

    assert replace.document.paragraphs[0].runs[0].bold
    assert replace.document.paragraphs[1].runs[0].bold
    assert replace.document.sections[0].footer.paragraphs[0].runs[0].bold
    assert replace.document.sections[0].header.paragraphs[0].runs[0].bold


def test_find_replace_overflow(
    data_dir: pathlib.Path,
) -> None:
    """Tests that find and replace does not remove additional text."""
    document = docx.Document(data_dir / "find_replace_test.docx")
    replace = writer.DocxReplace(document)

    replace.replace("This is a paragraph", "paragraph_replace")

    assert (
        replace.document.paragraphs[1].text
        == "Paragraph: paragraph_replace with overflow."
    )
