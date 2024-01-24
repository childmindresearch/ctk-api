"""Tests the diagnoses endpoints."""
import pathlib

import docx
import pytest
from fastapi import status, testclient

from . import conftest


@pytest.mark.asyncio()
async def test_post_md2docx(
    client: testclient.TestClient,
    endpoints: conftest.Endpoints,
    tmp_path: pathlib.Path,
) -> None:
    """Tests the anonymization endpoint."""
    md_text = "# Test Markdown Text"

    response = client.post(
        endpoints.POST_MARKDOWN_TO_DOCX,
        data={"markdown_text": md_text},
    )
    with (tmp_path / "md2docx.docx").open("wb") as output_file:
        output_file.write(response.content)
        document = docx.Document(tmp_path / "md2docx.docx")

    assert response.status_code == status.HTTP_200_OK
    assert (
        response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert (
        response.headers["Content-Disposition"] == 'attachment; filename="md2docx.docx"'
    )
    assert document.paragraphs[0].text == md_text[2:]
    assert document.paragraphs[0].style.name == "Heading 1"
