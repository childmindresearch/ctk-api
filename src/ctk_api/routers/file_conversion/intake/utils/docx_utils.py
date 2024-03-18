"""Utilities for the handling docx files."""

import re
from collections.abc import Iterable

import docx
from docx import oxml, table
from docx.enum import text
from docx.oxml import ns
from docx.text import paragraph as docx_paragraph


class DocxReplace:
    """Finds and replaces text in a Word document."""

    def __init__(self, document: docx.Document) -> None:
        """Initializes a DocxReplace object for finding and replacing.

        Args:
            document: The document to be written.

        """
        self.document = document

    def replace(self, find: str, replace: str) -> None:
        """Finds and replaces text in a Word document.

        Args:
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence..
        """
        for paragraph in self.document.paragraphs:
            self._replace_in_section(paragraph, find, replace)

        for section in self.document.sections:
            for paragraph in (
                *section.footer.paragraphs,
                *section.header.paragraphs,
            ):
                self._replace_in_section(paragraph, find, replace)

    @staticmethod
    def _replace_in_section(
        paragraph: docx_paragraph.Paragraph,
        find: str,
        replace: str,
    ) -> None:
        """Finds and replaces text in a Word document section.

        Applying a find and replace directly to the text property will lose
        formatting. python-docx splits each paragraph into runs, where each run
        has consistent styling. So we apply the find and replace to the runs.
        The replacement will gain the formatting of the first character of the
        text it replaces.

        This function modifies in place.

        Args:
            paragraph: The Word document's paragraph.
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence.

        """
        if find not in paragraph.text:
            return

        for paragraph_run in paragraph.runs:
            paragraph_run.text = paragraph_run.text.replace(find, replace)

        if find not in paragraph.text:
            return

        match_indices = [
            match.start() for match in re.finditer(re.escape(find), paragraph.text)
        ]

        run_lengths = [len(run.text) for run in paragraph.runs]
        for run_index in range(len(paragraph.runs) - 1, -1, -1):
            # Reverse loop so changes to run lengths don't affect the next
            # iteration.
            run_start = sum(run_lengths[:run_index])
            run_end = run_start + run_lengths[run_index]
            run_obj = paragraph.runs[run_index]
            for match_index in match_indices:
                if match_index < run_start or match_index >= run_end:
                    continue

                overflow = match_index + len(find) - run_end
                run_obj.text = run_obj.text[: match_index - run_start] + replace

                current_index = run_index + 1
                while overflow > 0:
                    next_run = paragraph.runs[current_index]
                    new_overflow = overflow - len(next_run.text)
                    next_run.text = next_run.text[overflow:]
                    overflow = new_overflow
                    current_index += 1

                break


def format_paragraphs(  # noqa: PLR0913
    paragraph: docx_paragraph.Paragraph | Iterable[docx_paragraph.Paragraph],
    *,
    line_spacing: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    bold: bool | None = None,
    italics: bool | None = None,
    font_size: int | None = None,
    font_rgb: tuple[int, int, int] | None = None,
    alignment: text.WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    """Formats a paragraph in a Word document.

    Args:
        paragraph: The paragraph to format.
        line_spacing: The line spacing of the paragraph.
        space_before: The spacing before the paragraph.
        space_after: The spacing after the paragraph.
        bold: Whether to bold the paragraph.
        italics: Whether to italicize the paragraph.
        italics: Whether to italicize the paragraph.
        font_size: The font size of the paragraph.
        font_rgb: The font color of the paragraph.
        alignment: The alignment of the paragraph.
    """
    if isinstance(paragraph, docx_paragraph.Paragraph):
        paragraph = [paragraph]

    for para in paragraph:
        if line_spacing is not None:
            para.paragraph_format.line_spacing = line_spacing

        if alignment is not None:
            para.alignment = alignment

        if space_before is not None:
            para.paragraph_format.space_before = space_before

        if space_after is not None:
            para.paragraph_format.space_after = space_after

        for run in para.runs:
            format_run(
                run,
                bold=bold,
                italics=italics,
                font_size=font_size,
                font_rgb=font_rgb,
            )


def format_run(
    run: docx.text.run.Run,
    *,
    bold: bool | None = None,
    italics: bool | None = None,
    font_size: int | None = None,
    font_rgb: tuple[int, int, int] | None = None,
) -> None:
    """Formats a run in a Word document.

    Args:
        run: The run to format.
        bold: Whether to bold the run.
        italics: Whether to italicize the run.
        font_size: The font size of the run.
        font_rgb: The font color of the run.
    """
    if bold is not None:
        run.bold = bold
    if italics is not None:
        run.italic = italics
    if font_size is not None:
        run.font.size = font_size
    if font_rgb is not None:
        run.font.color.rgb = docx.shared.RGBColor(*font_rgb)


def format_cell(  # noqa: PLR0913, D417
    cell: table._Cell,
    *,
    line_spacing: float | None = None,
    space_before: float | None = None,
    space_after: float | None = None,
    bold: bool | None = None,
    italics: bool | None = None,
    font_size: int | None = None,
    font_rgb: tuple[int, int, int] | None = None,
    background_rgb: tuple[int, int, int] | None = None,
    alignment: text.WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    """Formats a cell in a Word table.

    Args:
        cell: The cell to format.
        line_spacing: The line spacing of the cell.
        space_before: The spacing before the cell.
        space_after: The spacing after the cell.
        bold: Whether to bold the cell.
        italics: Whether to italicize the cell.
        font_size: The font size of the cell.
        font_rgb: The font color of the cell.
        background_rgb: The background color of the cell.
    """
    for para in cell.paragraphs:
        format_paragraphs(
            para,
            line_spacing=line_spacing,
            bold=bold,
            italics=italics,
            font_size=font_size,
            font_rgb=font_rgb,
            alignment=alignment,
            space_after=space_after,
            space_before=space_before,
        )

    if background_rgb is not None:
        shading = oxml.parse_xml(
            (r'<w:shd {} w:fill="' + f"{rgb_to_hex(*background_rgb)}" + r'"/>').format(
                ns.nsdecls("w"),
            ),
        )

        cell._tc.get_or_add_tcPr().append(shading)  # noqa: SLF001


def rgb_to_hex(r: int, g: int, b: int) -> str:
    """Converts RGB values to a hexadecimal color code.

    Args:
        r: The red component of the RGB color.
        g: The green component of the RGB color.
        b: The blue component of the RGB color.

    Returns:
        The hexadecimal color code representing the RGB color.
    """
    return f"#{r:02x}{g:02x}{b:02x}".upper()
