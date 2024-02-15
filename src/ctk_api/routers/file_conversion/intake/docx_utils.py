"""Utilities for handling DOCX files."""
import dataclasses
import re

import docx
import mlconjug3
import spacy
from spacy import symbols, tokens

NLP = spacy.load("en_core_web_sm")


@dataclasses.dataclass
class SubjectVerbPair:
    """A pair of a subject and a verb in a sentence."""

    sentence: str
    subject_token: tokens.Token
    verb_token: tokens.Token

    @property
    def subject(self) -> str:
        """The subject of the sentence."""
        return self.subject_token.text

    @property
    def verb(self) -> str:
        """The verb of the sentence."""
        return self.verb_token.text

    @property
    def verb_indices(self) -> tuple[int, int]:
        """The indices of the verb in the sentence."""
        return self.verb_token.idx, self.verb_token.idx + len(self.verb)


class DocumentCorrections:
    """Corrects verb conjugations associated with 'they' in a Word document."""

    def __init__(
        self,
        document: docx.Document,
        *,
        correct_they: bool = True,
        correct_capitalization: bool = True,
    ) -> None:
        """Initializes the corrector with a document.

        Args:
            document: The docx document to correct.
            correct_they: Whether to correct verb conjugations associated with 'they'.
            correct_capitalization: Whether to correct the capitalization of the
                first word.
        """
        self.document = document
        self.replacer = DocxReplace(document)
        self.conjugator = mlconjug3.Conjugator(language="en")
        self.correct_they = correct_they
        self.correct_capitalization = correct_capitalization

    def correct(self) -> None:
        """Corrects verb conjugations associated with 'they' in the document."""
        if not self.correct_they and not self.correct_capitalization:
            return

        for paragraph in self.document.paragraphs:
            self._correct_paragraph(paragraph)

    def _correct_paragraph(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Corrects conjugations in a single paragraph.

        Args:
            paragraph: The paragraph to correct.
        """
        sentences = NLP(paragraph.text).sents
        for sentence in sentences:
            self._correct_sentence(sentence.text)

    def _correct_sentence(self, sentence: str) -> None:
        """Corrects the conjugation of verbs associated with they in a sentence.

        Args:
            sentence: The sentence to correct.
        """
        words = sentence.split(" ")
        if "they" not in [word.lower() for word in words]:
            return

        subject_verb_pairs = self._find_subject_verb(sentence)
        they_verb_pairs = [
            pair for pair in subject_verb_pairs if pair.subject.lower() == "they"
        ]

        for pair in reversed(they_verb_pairs):
            sentence = self._correct_verb_conjugation(sentence, pair)

        if sentence[0].islower():
            new_sentence = sentence[0].upper() + sentence[1:]
            self.replacer.replace(sentence, new_sentence)

    def _correct_verb_conjugation(
        self,
        sentence: str,
        pair: SubjectVerbPair,
    ) -> str:
        """Corrects the verb conjugation associated with 'they' in a sentence.

        Args:
            sentence: The sentence to correct.
            pair: The pair of subject and verb in the sentence.
        """
        for child in pair.verb_token.children:
            if child.tag_ == "VBZ" and child.lemma_ in ["be", "have"]:
                pair.verb_token = child
            elif child.tag_ == "VBZ":
                child_pair = SubjectVerbPair(
                    sentence=sentence,
                    subject_token=pair.subject_token,
                    verb_token=child,
                )
                sentence = self._correct_verb_conjugation(sentence, child_pair)
        if pair.verb_token.tag_ != "VBZ":
            return sentence

        verb = self.conjugator.conjugate(pair.verb_token.lemma_)
        if isinstance(verb, list):
            verb = verb[0]
        try:
            conjugated_verb = verb["indicative"]["indicative present"]["they"]  # type: ignore[index]
        except TypeError as exc_info:
            if "'NoneType' object is not subscriptable" in str(exc_info):
                # Verb is unknown to the conjugator.
                return sentence
            raise

        new_sentence = (
            sentence[: pair.verb_indices[0]]
            + conjugated_verb
            + sentence[pair.verb_indices[1] :]
        )

        self.replacer.replace(sentence, new_sentence)
        return new_sentence

    @staticmethod
    def _find_subject_verb(sentence: str) -> list[SubjectVerbPair]:
        """Finds the subject and verb of a sentence, only if the subject.

        Args:
            sentence: The sentence to analyze.

        Returns:
            A list of indices of subjects and verbs
        """
        doc = NLP(sentence)
        return [
            SubjectVerbPair(
                sentence=sentence,
                subject_token=word,
                verb_token=word.head,
            )
            for word in doc
            if word.dep in [symbols.nsubj, symbols.nsubjpass]
            and word.head.tag_.startswith("VB")
        ]


class DocxReplace:
    """Finds and replaces text in a Word document."""

    def __init__(self, document: docx.Document) -> None:
        """Initializes a DocxReplace object for finding and replacing.

        Args:
            document: The document to be written.

        """
        self.document = document

    def replace(self, find: str, replace: str) -> None:
        """Finds and replaces text in a Word document.

        Args:
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence..
        """
        for paragraph in self.document.paragraphs:
            self._replace_in_section(paragraph, find, replace)

        for section in self.document.sections:
            for paragraph in (
                *section.footer.paragraphs,
                *section.header.paragraphs,
            ):
                self._replace_in_section(paragraph, find, replace)

    @staticmethod
    def _replace_in_section(
        paragraph: docx.text.paragraph.Paragraph,
        find: str,
        replace: str,
    ) -> None:
        """Finds and replaces text in a Word document section.

        Applying a find and replace directly to the text property will lose
        formatting. python-docx splits each paragraph into runs, where each run
        has consistent styling. So we apply the find and replace to the runs.
        The replacement will gain the formatting of the first character of the
        text it replaces.

        This function modifies in place.

        Args:
            paragraph: The Word document's paragraph.
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence.

        """
        if find not in paragraph.text:
            return

        for paragraph_run in paragraph.runs:
            paragraph_run.text = paragraph_run.text.replace(find, replace)

        if find not in paragraph.text:
            return

        match_indices = [
            match.start() for match in re.finditer(re.escape(find), paragraph.text)
        ]

        run_lengths = [len(run.text) for run in paragraph.runs]
        for run_index in range(len(paragraph.runs) - 1, -1, -1):
            # Reverse loop so changes to run lengths don't affect the next
            # iteration.
            run_start = sum(run_lengths[:run_index])
            run_end = run_start + run_lengths[run_index]
            run_obj = paragraph.runs[run_index]
            for match_index in match_indices:
                if match_index < run_start or match_index >= run_end:
                    continue

                overflow = match_index + len(find) - run_end
                run_obj.text = run_obj.text[: match_index - run_start] + replace

                current_index = run_index + 1
                while overflow > 0:
                    next_run = paragraph.runs[current_index]
                    new_overflow = overflow - len(next_run.text)
                    next_run.text = next_run.text[overflow:]
                    overflow = new_overflow
                    current_index += 1

                break
