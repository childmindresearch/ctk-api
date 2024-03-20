"""Contains report writing functionality for intake information."""

import enum
import itertools
from typing import ParamSpec, TypeVar

import docx
from docx.enum import table as enum_table
from docx.enum import text as enum_text
from docx.text import paragraph as docx_paragraph

from ctk_api.core import config
from ctk_api.routers.file_conversion.intake import descriptors, parser
from ctk_api.routers.file_conversion.intake.utils import (
    docx_utils,
    language_utils,
    string_utils,
)

P = ParamSpec("P")
T = TypeVar("T")

DATA_DIR = config.DATA_DIR
RGB_INTAKE = (178, 161, 199)
RGB_TESTING = (155, 187, 89)
RGB_TEMPLATE = (247, 150, 70)
PLACEHOLDER = "______"


class Style(enum.Enum):
    """The styles for the report."""

    HEADING_1 = "Heading 1"
    HEADING_2 = "Heading 2"
    HEADING_3 = "Heading 3"
    TITLE = "Title"
    NORMAL = "Normal"


class ReportWriter:
    """Writes a report for intake information."""

    def __init__(self, intake: parser.IntakeInformation) -> None:
        """Initializes the report writer.

        Args:
            intake: The intake information.
        """
        self.intake = intake
        self.report = docx.Document(DATA_DIR / "report_template.docx")
        self.insert_before = next(
            paragraph
            for paragraph in self.report.paragraphs
            if "MENTAL STATUS EXAMINATION AND TESTING BEHAVIORAL OBSERVATIONS"
            in paragraph.text
        )

        if not self.insert_before:
            msg = "Insertion point not found in the report template."
            raise ValueError(msg)

    def transform(self) -> None:
        """Transforms the intake information to a report."""
        self.write_reason_for_visit()
        self.write_developmental_history()
        self.write_academic_history()
        self.write_social_history()
        self.write_psychiatric_history()
        self.write_medical_history()
        self.write_current_psychiatric_functioning()
        self.add_page_break()
        self.replace_patient_information()
        self.apply_corrections()
        self.add_signatures()

    def replace_patient_information(self) -> None:
        """Replaces the patient information in the report."""
        replacements = {
            "full_name": self.intake.patient.full_name,
            "preferred_name": self.intake.patient.first_name,
            "date_of_birth": self.intake.patient.date_of_birth.strftime("%m/%d/%Y"),
            "reporting_guardian": self.intake.patient.guardian.title_name,
            "aged_gender": self.intake.patient.age_gender_label,
            "pronoun_0": self.intake.patient.pronouns[0],
            "pronoun_1": self.intake.patient.pronouns[1],
            "pronoun_2": self.intake.patient.pronouns[2],
            "pronoun_4": self.intake.patient.pronouns[4],
        }

        for template, replacement in replacements.items():
            template_formatted = "{{" + template.upper() + "}}"
            docx_utils.DocxReplace(self.report).replace(template_formatted, replacement)

    def write_reason_for_visit(self) -> None:
        """Writes the reason for visit to the end of the report."""
        patient = self.intake.patient
        handedness = patient.handedness
        iep = patient.education.individualized_educational_program
        past_diagnoses = patient.psychiatric_history.past_diagnoses.transform(
            short=True,
        )
        classroom = patient.education.classroom_type

        if patient.education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(
                int(patient.education.grade),
            )
        else:
            grade_superscript = ""

        texts = [
            f"""
            At the time of enrollment, {patient.first_name} was a
            {patient.age}-year-old, {handedness} {patient.age_gender_label}
            {past_diagnoses}. {patient.first_name} was placed in a {classroom}
            {patient.education.grade}""",
            f"{grade_superscript} ",
            f"""
            grade classroom at {patient.education.school_name}.
            {patient.first_name} {iep}. {patient.first_name} and
            {patient.pronouns[2]} {patient.guardian.relationship},
            {patient.guardian.title_full_name}, attended the present evaluation due to
            concerns regarding {PLACEHOLDER}. The family is hoping for
            {PLACEHOLDER}. The family learned of the study through
            {PLACEHOLDER}.
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        heading = self._insert("REASON FOR VISIT", Style.HEADING_1)
        paragraph = self._insert(texts[0])
        paragraph.add_run(texts[1]).font.superscript = True
        paragraph.add_run(" " + texts[2])
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_developmental_history(self) -> None:
        """Writes the developmental history to the end of the report."""
        heading = self._insert("DEVELOPMENTAL HISTORY", Style.HEADING_1)
        docx_utils.format_paragraphs(heading, font_rgb=RGB_INTAKE)
        self.write_prenatal_history()
        self.write_developmental_milestones()
        self.write_early_education()

    def write_prenatal_history(self) -> None:
        """Writes the prenatal and birth history of the patient to the report."""
        patient = self.intake.patient
        development = patient.development
        pregnancy_symptoms = development.birth_complications
        delivery = development.delivery
        delivery_location = development.delivery_location
        adaptability = development.adaptability
        duration_of_pregnancy = development.duration_of_pregnancy

        text = f"""
            {patient.guardian.title_name} reported {pregnancy_symptoms}.
            {patient.first_name} was born at
            {duration_of_pregnancy} of gestation with {delivery} at
            {delivery_location}. {patient.first_name} had {adaptability}
            during infancy and was {development.soothing_difficulty.name} to
            soothe.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Prenatal and Birth History", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_developmental_milestones(self) -> None:
        """Writes the developmental milestones to the report."""
        patient = self.intake.patient
        started_walking = patient.development.started_walking
        started_talking = patient.development.started_talking
        daytime_dryness = patient.development.daytime_dryness
        nighttime_dryness = patient.development.nighttime_dryness

        text = f"""
            {patient.first_name}'s achievement of social, language, fine and
            gross motor developmental milestones were within normal limits, as
            reported by {patient.guardian.title_name}. {patient.first_name}
            {started_walking} and {started_talking}.
            {patient.pronouns[0].capitalize()} {daytime_dryness} and
            {nighttime_dryness}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Developmental Milestones", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_early_education(self) -> None:
        """Writes the early education information to the report."""
        patient = self.intake.patient
        development = patient.development

        reporting_guardian = patient.guardian.title_name
        early_intervention = development.early_intervention_age
        cpse = development.cpse_age

        text = f"""
            {reporting_guardian} reported that
            {patient.first_name} {early_intervention} and {cpse}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Early Educational Interventions", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_academic_history(self) -> None:
        """Writes the academic history to the end of the report."""
        heading = self._insert("ACADEMIC AND EDUCATIONAL HISTORY", Style.HEADING_1)
        docx_utils.format_paragraphs(heading, font_rgb=RGB_INTAKE)
        self.write_previous_testing()
        self.write_academic_history_table()
        self.write_educational_history()

    def write_previous_testing(self) -> None:
        """Writes the previous testing information to the report."""
        patient = self.intake.patient

        text = f"""
        {patient.first_name} has no history of previous psychoeducational
        evaluations./{patient.first_name} was evaluated by {PLACEHOLDER} in 20XX.
        Documentation of the results of the evaluation(s) were unavailable at
        the time of writing this report/ Notable results include:
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Previous Testing", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_TEMPLATE)

    def write_academic_history_table(self) -> None:
        """Writes the academic history table to the report."""
        paragraph = self._insert("Name, Date of Assessment")
        docx_utils.format_paragraphs(
            paragraph,
            font_rgb=RGB_INTAKE,
            bold=True,
            alignment=enum_text.WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

        table = self.report.add_table(7, 4)
        self.insert_before._p.addprevious(table._tbl)  # noqa: SLF001
        table.style = "Table Grid"
        header_row = table.rows[0].cells

        header_texts = [
            "Domain/Index/Subtest",
            "Standard Score",
            "Percentile Rank",
            "Descriptor",
        ]
        for i, header in enumerate(header_texts):
            header_row[i].text = header
            header_row[i].width = 10
            docx_utils.format_cell(
                header_row[i],
                bold=True,
                font_rgb=RGB_INTAKE,
                background_rgb=(217, 217, 217),
                alignment=enum_text.WD_ALIGN_PARAGRAPH.CENTER,
            )
        for row in table.rows:
            row.height = 1
            row.height_rule = enum_table.WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                docx_utils.format_cell(
                    cell,
                    line_spacing=1,
                    space_after=0,
                    space_before=0,
                )

    def write_educational_history(self) -> None:
        """Writes the educational history to the report."""
        patient = self.intake.patient
        education = patient.education
        has_iep = (
            education.individualized_educational_program.base
            == descriptors.IndividualizedEducationProgram.yes.value
        )

        if has_iep:
            iep_prior_text = f"""
                {patient.first_name} was
                granted an Individualized Education Program (IEP) in
                {PLACEHOLDER} grade due to {PLACEHOLDER}
                difficulties.
            """
        else:
            iep_prior_text = f"""{patient.first_name} has never had an
                              Individiualized Education Program (IEP)."""
        if education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(education.grade)
        else:
            grade_superscript = ""
        past_schools = education.past_schools

        text_prior = f"""
            {patient.first_name} {past_schools}. {patient.pronouns[0]}
            previously struggled with (provide details of academic challenges
            and behavioral difficulties in school). {iep_prior_text}
        """
        texts_current = [
            f"""{patient.first_name} is currently in the {education.grade}""",
            f"{grade_superscript} ",
            f"""
                grade at {education.school_name}.
                {patient.first_name} does/does not receive special
                education services and maintains/does not have an IEP
                allowing accommodations for/including {PLACEHOLDER}.
                {patient.first_name} is generally an average/above
                average/below average student and receives mostly (describe
                grades). [Describe any academic issues reported by parent or
                child.] {patient.first_name} continues to exhibit
                weaknesses in {PLACEHOLDER}.
            """,
        ]
        text_prior = string_utils.remove_excess_whitespace(text_prior)
        texts_current = [
            string_utils.remove_excess_whitespace(text) for text in texts_current
        ]

        heading = self._insert("Educational History", Style.HEADING_2)
        prior_paragraph = self._insert(text_prior)
        docx_utils.format_paragraphs((heading, prior_paragraph), font_rgb=RGB_INTAKE)

        current_paragraph = self._insert(texts_current[0])
        current_paragraph.add_run(texts_current[1]).font.superscript = True
        current_paragraph.add_run(" " + texts_current[2])
        docx_utils.format_paragraphs(current_paragraph, font_rgb=RGB_TEMPLATE)

    def write_social_history(self) -> None:
        """Writes the social history to the end of the report."""
        heading = self._insert("SOCIAL HISTORY", Style.HEADING_1)
        docx_utils.format_paragraphs(heading, font_rgb=RGB_INTAKE)
        self.write_home_and_adaptive_functioning()
        self.write_social_functioning()

    def write_home_and_adaptive_functioning(self) -> None:
        """Writes the home and adaptive functioning to the report."""
        patient = self.intake.patient
        household = patient.household
        language_fluencies = self._join_patient_languages(self.intake.patient.languages)

        text_home = f"""
            {patient.first_name} lives in {household.city},
            {household.state}, with {household.members}.
            The {patient.guardian.parent_or_guardian}s are
            {household.guardian_marital_status}.
            {string_utils.join_with_oxford_comma(household.languages)} {"are" if
            len(household.languages) > 1 else "is"} spoken at home.
            {patient.language_spoken_best} is reportedly
            {patient.first_name}'s preferred language.
            {patient.first_name} {language_fluencies}.
        """

        text_adaptive = f"""
            {patient.guardian.title_name} denied any concerns with {patient.pronouns[2]}
            functioning in the home setting// Per {patient.guardian.title_name},
            {patient.first_name} has a history of {PLACEHOLDER} (temper
            outbursts, oppositional behaviors, etc.) in the home setting. (Write
            details of behavioral difficulties). (Also include any history of
            sleep difficulties, daily living skills, poor hygiene, etc.)
            """
        text_home = string_utils.remove_excess_whitespace(text_home)
        text_adaptive = string_utils.remove_excess_whitespace(text_adaptive)

        heading = self._insert("Home and Adaptive Functioning", Style.HEADING_2)
        home_paragraph = self._insert(text_home)
        adaptive_paragraph = self._insert(text_adaptive)

        docx_utils.format_paragraphs((heading, home_paragraph), font_rgb=RGB_INTAKE)
        docx_utils.format_paragraphs(adaptive_paragraph, font_rgb=RGB_TEMPLATE)

    def write_social_functioning(self) -> None:
        """Writes the social functioning to the report."""
        patient = self.intake.patient

        text = f"""
            {patient.guardian.title_name} was pleased to describe
            {patient.first_name} as a (insert adjective e.g., affectionate)
            {patient.age_gender_label}. {patient.guardian.title_name} reported
            that {patient.pronouns[0]} has many/several/one friends in
            {patient.pronouns[2]} peer group in school and on
            {patient.pronouns[2]} team/club/etc. {patient.first_name}
            socializes with friends outside of school and has a
            (positive/fair/poor) relationship with them.
            {patient.first_name}'s hobbies include {PLACEHOLDER}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Social Functioning", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_psychiatric_history(self) -> None:
        """Writes the psychiatric history to the end of the report."""
        heading = self._insert("PSYCHRIATIC HISTORY", Style.HEADING_1)
        docx_utils.format_paragraphs(heading, font_rgb=RGB_INTAKE)
        self.write_past_psychriatic_diagnoses()
        self.write_past_psychiatric_hospitalizations()
        self.write_past_therapeutic_interventions()
        self.write_past_self_injurious_behaviors_and_suicidality()
        self.write_past_aggressive_behaviors_and_homicidality()
        self.expose_to_violence_and_trauma()
        self.administration_for_childrens_services_involvement()
        self.write_family_psychiatric_history()

    def write_past_psychiatric_hospitalizations(self) -> None:
        """Writes the past psychiatric hospitalizations to the report."""
        patient = self.intake.patient
        text = f"""
            {patient.guardian.title_name} denied any history of past psychiatric
            hospitalizations for {patient.first_name}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Past Psychiatric Hospitalizations", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_TEMPLATE)

    def administration_for_childrens_services_involvement(self) -> None:
        """Writes the ACS involvement to the report."""
        patient = self.intake.patient

        text = str(patient.psychiatric_history.children_services)
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert(
            "Administration for Children's Services (ACS) Involvement",
            Style.HEADING_2,
        )
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_past_aggressive_behaviors_and_homicidality(self) -> None:
        """Writes the past aggressive behaviors and homicidality to the report."""
        patient = self.intake.patient

        text = str(patient.psychiatric_history.aggresive_behaviors)
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert(
            "Past Severe Aggressive Behaviors and Homicidality",
            Style.HEADING_2,
        )
        report = self._insert(text)
        docx_utils.format_paragraphs((heading, report), font_rgb=RGB_INTAKE)

    def write_past_psychriatic_diagnoses(self) -> None:
        """Writes the past psychiatric diagnoses to the report."""
        patient = self.intake.patient
        past_diagnoses = patient.psychiatric_history.past_diagnoses.transform(
            short=False,
        )

        text = f"""
            {patient.first_name} {past_diagnoses}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Past Psychiatric Diagnoses", Style.HEADING_2)
        report = self._insert(text)
        docx_utils.format_paragraphs((heading, report), font_rgb=RGB_INTAKE)

    def write_family_psychiatric_history(self) -> None:
        """Writes the family psychiatric history to the report."""
        patient = self.intake.patient
        text = f"""
        {patient.first_name}'s family history is largely unremarkable for
        psychiatric illnesses. {patient.guardian.title_name} denied any family
        history related to homicidality, suicidality, depression, bipolar
        disorder, attention-deficit/hyperactivity disorder, autism spectrum
        disorder, learning disorders, psychotic disorders, eating disorders,
        oppositional defiant or conduct disorders, substance abuse, panic,
        generalized anxiety, or obsessive-compulsive disorders. Information
        regarding {patient.first_name}'s family psychiatric history was
        deferred."""
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Family Psychiatric History", Style.HEADING_2)
        report = self._insert(text)
        docx_utils.format_paragraphs((heading, report), font_rgb=RGB_TEMPLATE)

    def write_past_therapeutic_interventions(self) -> None:
        """Writes the past therapeutic history to the report."""
        patient = self.intake.patient
        guardian = patient.guardian
        interventions = patient.psychiatric_history.therapeutic_interventions

        if not interventions:
            texts = [
                f"""
                    {patient.guardian.title_name} denied any history of therapeutic
                    interventions.
                """,
            ]
        else:
            texts = [
                f"""
                    From {intervention.start}-{intervention.end},
                    {patient.first_name} engaged in therapy with
                    {intervention.therapist} due to "{intervention.reason}" at a
                    frequency of {intervention.frequency}. {guardian.title_name}
                    described the treatment as "{intervention.effectiveness}".
                    Treatment was ended due to "{intervention.reason_ended}".
                    """
                for intervention in interventions
            ]

        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        heading = self._insert("Past Therapeutic Interventions", Style.HEADING_2)
        paragraphs = [self._insert(text) for text in texts]
        docx_utils.format_paragraphs((heading, *paragraphs), font_rgb=RGB_INTAKE)

    def write_past_self_injurious_behaviors_and_suicidality(self) -> None:
        """Writes the past self-injurious behaviors and suicidality to the report."""
        patient = self.intake.patient

        text = str(patient.psychiatric_history.self_harm)
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert(
            "Past Self-Injurious Behaviors and Suicidality",
            Style.HEADING_2,
        )
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def expose_to_violence_and_trauma(self) -> None:
        """Writes the exposure to violence and trauma to the report."""
        patient = self.intake.patient

        text = str(patient.psychiatric_history.violence_and_trauma)
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Exposure to Violence and Trauma", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_medical_history(self) -> None:
        """Writes the medical history to the end of the report."""
        patient = self.intake.patient

        text = f"""
            {patient.first_name}'s medical history is unremarkable for
            significant medical conditions. {patient.pronouns[0]} is not
            currently taking any medications for chronic medical conditions.
            {patient.first_name} wears prescription glasses in home and
            school settings. {patient.pronouns[0]} does/does not require a
            hearing device. {patient.guardian.title_name} denied any history of
            seizures, head trauma, migraines, meningitis or encephalitis.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("MEDICAL HISTORY", Style.HEADING_1)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_TEMPLATE)

    def write_clinical_summary_and_impressions(self) -> None:
        """Writes the clinical summary and impressions to the report."""
        patient = self.intake.patient
        gender = patient.age_gender_label

        text = f"""
            {patient.first_name} is a
            sociable/resourceful/pleasant/hardworking/etc. {gender} who
            participated in the Healthy Brain Network research project through
            the Child Mind Institute in the interest of participating in
            research/due to parental concerns regarding {PLACEHOLDER}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("CLINICAL SUMMARY AND IMPRESSIONS", Style.HEADING_1)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_TESTING)

    def write_current_psychiatric_functioning(self) -> None:
        """Writes the current psychiatric functioning to the report.

        Note: this section mixes color codings. Color decorators are applied
        to the called functions instead.
        """
        heading = self._insert("CURRENT PSYCHIATRIC FUNCTIONING", Style.HEADING_1)
        docx_utils.format_paragraphs(heading, font_rgb=RGB_INTAKE)
        self.write_current_psychiatric_medications_intake()
        self.write_current_psychiatric_medications_testing()
        self.write_denied_symptoms()

    def write_current_psychiatric_medications_intake(self) -> None:
        """Writes the current psychiatric medications to the report."""
        patient = self.intake.patient
        text = f"""
        {patient.first_name} is currently prescribed a daily/twice daily
        oral course of {PLACEHOLDER} for {PLACEHOLDER}. {patient.pronouns[0]} is
        being treated by Doctortype, DoctorName, monthly/weekly/biweekly. The
        medication has been ineffective/effective.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("Current Psychiatric Medications", Style.HEADING_2)
        paragraph = self._insert(text)
        docx_utils.format_paragraphs((heading, paragraph), font_rgb=RGB_INTAKE)

    def write_current_psychiatric_medications_testing(self) -> None:
        """Writes the current psychiatric medications to the report."""
        patient = self.intake.patient
        texts = [
            f"""
        [Rule out presenting diagnoses, using headlines and KSADS/DSM criteria.
        Examples of headlines include: Temper Outbursts (ending should include
        “{patient.guardian.title_name} denied any consistent patterns of
        irritability for {patient.first_name}" if applicable), Inattention
        and Hyperactivity, Autism-Related Symptoms, Oppositional Defiant
        Behaviors, etc.].""",
            "Establish a baseline first for temper outbursts",
            f"""
           (Ex:
        Though {patient.first_name} is generally a {PLACEHOLDER} child,
        {patient.pronouns[0]} continues to have difficulties with temper
        tantrums…).
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        paragraph = self._insert(texts[0])
        paragraph.add_run(texts[1])
        paragraph.runs[-1].bold = True
        paragraph.add_run(texts[2])
        docx_utils.format_paragraphs(paragraph, italics=True, font_rgb=RGB_TESTING)

    def write_denied_symptoms(self) -> None:
        """Writes the denied symptoms to the report."""
        patient = self.intake.patient
        text = f"""
        {patient.guardian.title_name} and {patient.first_name} denied any
        current significant symptoms related to mood, suicidality, psychosis,
        eating, oppositional or conduct behaviors, substance abuse, autism,
        tics, inattention/hyperactivity, enuresis/encopresis, trauma, sleep,
        panic, anxiety or obsessive-compulsive disorders.
        """
        text = string_utils.remove_excess_whitespace(text)

        paragraph = self._insert(text)
        docx_utils.format_paragraphs(paragraph, font_rgb=RGB_TESTING)

    def apply_corrections(self) -> None:
        """Applies various grammatical and styling corrections."""
        document_corrector = language_utils.DocumentCorrections(
            self.report,
            correct_they=self.intake.patient.pronouns[0] == "they",
            correct_capitalization=True,
        )
        document_corrector.correct()

    def add_signatures(self) -> None:
        """Adds the signatures to the report.

        Michael Milham's signature is placed in a different location than the
        other signatures. As such, it needs some custom handling.
        """
        signature_dir = DATA_DIR / "signatures"
        for signature in signature_dir.glob("*.png"):
            name = signature.stem.replace("_", " ")
            paragraph_index = next(
                index
                for index in range(len(self.report.paragraphs))
                if self.report.paragraphs[index].text.lower().startswith(name)
            )

            if name != "michael p. milham":
                # All signatures except for Michael P. Milham's are inserted
                # on above a underlined paragraph one before the one that contains
                # the name.
                paragraph_index -= 1

            image_paragraph = docx_utils.insert_image_before(
                self.report.paragraphs[paragraph_index],
                signature,
            )
            if name != "michael p. milham":
                docx_utils.insert_paragraph_before(image_paragraph, "")

    def add_page_break(self) -> None:
        """Adds a page break to the report."""
        paragraph = self._insert("")
        paragraph.add_run().add_break(enum_text.WD_BREAK.PAGE)

    def _insert(
        self,
        text: str,
        style: Style = Style.NORMAL,
    ) -> docx_paragraph.Paragraph:
        """Inserts text at the insertion point.

        Given the current structure of the report, text insertions only occur
        in one location.

        Args:
            text: The text to insert.
            style: The style of the text.

        Returns:
            The new paragraph.
        """
        return docx_utils.insert_paragraph_before(
            self.insert_before,
            text,
            style=style.value,
        )

    @staticmethod
    def _join_patient_languages(languages: list[parser.Language]) -> str:
        """Joins the patient's languages."""
        fluency_groups = itertools.groupby(
            languages,
            key=lambda language: language.fluency,
        )
        fluency_dict = {
            fluency: [language.name for language in language_group]
            for fluency, language_group in fluency_groups
        }

        language_descriptions = [
            f"{fluency} in {string_utils.join_with_oxford_comma(fluency_dict[fluency])}"
            for fluency in ["fluent", "proficient", "conversational"]
            if fluency in fluency_dict
        ]
        prepend_is = len(language_descriptions) > 0
        if "basic" in fluency_dict:
            language_descriptions.append(
                (
                    "has basic skills in "
                    f"{string_utils.join_with_oxford_comma(fluency_dict['basic'])}"
                ),
            )

        text = string_utils.join_with_oxford_comma(language_descriptions)
        if prepend_is:
            text = "is " + text
        return text
