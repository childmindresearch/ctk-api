"""Utilities for the intake parser."""
import dataclasses
import re

import docx
import fastapi
import mlconjug3
import polars as pl
import spacy
from docx import oxml, table
from docx.enum import text
from docx.oxml import ns
from docx.text import paragraph as docx_paragraph
from fastapi import status
from spacy import symbols, tokens

NLP = spacy.load("en_core_web_sm")


@dataclasses.dataclass
class SubjectVerbPair:
    """A pair of a subject and a verb in a sentence."""

    sentence: str
    subject_token: tokens.Token
    verb_token: tokens.Token

    @property
    def subject(self) -> str:
        """The subject of the sentence."""
        return self.subject_token.text

    @property
    def verb(self) -> str:
        """The verb of the sentence."""
        return self.verb_token.text

    @property
    def verb_indices(self) -> tuple[int, int]:
        """The indices of the verb in the sentence."""
        return self.verb_token.idx, self.verb_token.idx + len(self.verb)


class DocumentCorrections:
    """Corrects verb conjugations associated with 'they' in a Word document."""

    def __init__(
        self,
        document: docx.Document,
        *,
        correct_they: bool = True,
        correct_capitalization: bool = True,
    ) -> None:
        """Initializes the corrector with a document.

        Args:
            document: The docx document to correct.
            correct_they: Whether to correct verb conjugations associated with 'they'.
            correct_capitalization: Whether to correct the capitalization of the
                first word.
        """
        self.document = document
        self.replacer = DocxReplace(document)
        self.conjugator = mlconjug3.Conjugator(language="en")
        self.correct_they = correct_they
        self.correct_capitalization = correct_capitalization

    def correct(self) -> None:
        """Corrects verb conjugations associated with 'they' in the document."""
        if not self.correct_they and not self.correct_capitalization:
            return

        for paragraph in self.document.paragraphs:
            self._correct_paragraph(paragraph)

    def _correct_paragraph(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Corrects conjugations in a single paragraph.

        Args:
            paragraph: The paragraph to correct.
        """
        sentences = NLP(paragraph.text).sents
        for sentence in sentences:
            self._correct_sentence(sentence.text)

    def _correct_sentence(self, sentence: str) -> None:
        """Corrects the conjugation of verbs associated with they in a sentence.

        Args:
            sentence: The sentence to correct.
        """
        words = sentence.split(" ")
        if "they" not in [word.lower() for word in words]:
            return

        subject_verb_pairs = self._find_subject_verb(sentence)
        they_verb_pairs = [
            pair for pair in subject_verb_pairs if pair.subject.lower() == "they"
        ]

        for pair in reversed(they_verb_pairs):
            sentence = self._correct_they_verb_conjugation(sentence, pair)

        if sentence[0].islower():
            new_sentence = sentence[0].upper() + sentence[1:]
            self.replacer.replace(sentence, new_sentence)

    def _correct_they_verb_conjugation(
        self,
        sentence: str,
        pair: SubjectVerbPair,
    ) -> str:
        """Corrects the verb conjugation associated with 'they' in a sentence.

        Args:
            sentence: The sentence to correct.
            pair: The pair of subject and verb in the sentence.

        Notes:
            To my knowledge, 'to be' is the only verb that has a different
            conjugation for third person singular/plural in the past tense. If
            other verbs are found to have different conjugations in the past
            tense, this function will need to be updated.
        """
        for child in pair.verb_token.children:
            if child.tag_ == "VBZ" and child.lemma_ in ["be", "have"]:
                pair.verb_token = child
            elif child.tag_ == "VBZ":
                child_pair = SubjectVerbPair(
                    sentence=sentence,
                    subject_token=pair.subject_token,
                    verb_token=child,
                )
                sentence = self._correct_they_verb_conjugation(sentence, child_pair)
        if pair.verb_token.tag_ != "VBZ" and not (
            pair.verb_token.tag_ == "VBD" and pair.verb_token.lemma_ == "be"
        ):
            return sentence

        verb = self.conjugator.conjugate(pair.verb_token.lemma_)
        if isinstance(verb, list):
            verb = verb[0]
        try:
            if pair.verb_token.tag_ == "VBZ":
                conjugated_verb = verb["indicative"]["indicative present"]["they"]  # type: ignore[index]
            else:
                conjugated_verb = verb["indicative"]["indicative past tense"]["they"]  # type: ignore[index]
        except TypeError as exc_info:
            if "'NoneType' object is not subscriptable" in str(exc_info):
                # Verb is unknown to the conjugator.
                return sentence
            raise

        new_sentence = (
            sentence[: pair.verb_indices[0]]
            + conjugated_verb
            + sentence[pair.verb_indices[1] :]
        )

        self.replacer.replace(sentence, new_sentence)
        return new_sentence

    @staticmethod
    def _find_subject_verb(sentence: str) -> list[SubjectVerbPair]:
        """Finds the subject and verb of a sentence, only if the subject.

        Args:
            sentence: The sentence to analyze.

        Returns:
            A list of indices of subjects and verbs
        """
        doc = NLP(sentence)
        return [
            SubjectVerbPair(
                sentence=sentence,
                subject_token=word,
                verb_token=word.head,
            )
            for word in doc
            if word.dep in [symbols.nsubj, symbols.nsubjpass]
            and word.head.tag_.startswith("VB")
        ]


class DocxReplace:
    """Finds and replaces text in a Word document."""

    def __init__(self, document: docx.Document) -> None:
        """Initializes a DocxReplace object for finding and replacing.

        Args:
            document: The document to be written.

        """
        self.document = document

    def replace(self, find: str, replace: str) -> None:
        """Finds and replaces text in a Word document.

        Args:
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence..
        """
        for paragraph in self.document.paragraphs:
            self._replace_in_section(paragraph, find, replace)

        for section in self.document.sections:
            for paragraph in (
                *section.footer.paragraphs,
                *section.header.paragraphs,
            ):
                self._replace_in_section(paragraph, find, replace)

    @staticmethod
    def _replace_in_section(
        paragraph: docx_paragraph.Paragraph,
        find: str,
        replace: str,
    ) -> None:
        """Finds and replaces text in a Word document section.

        Applying a find and replace directly to the text property will lose
        formatting. python-docx splits each paragraph into runs, where each run
        has consistent styling. So we apply the find and replace to the runs.
        The replacement will gain the formatting of the first character of the
        text it replaces.

        This function modifies in place.

        Args:
            paragraph: The Word document's paragraph.
            find: The text to find.
            replace: The text to replace.
            capitalize: Whether to capitalize the replacement if it is the first
                word of a sentence.

        """
        if find not in paragraph.text:
            return

        for paragraph_run in paragraph.runs:
            paragraph_run.text = paragraph_run.text.replace(find, replace)

        if find not in paragraph.text:
            return

        match_indices = [
            match.start() for match in re.finditer(re.escape(find), paragraph.text)
        ]

        run_lengths = [len(run.text) for run in paragraph.runs]
        for run_index in range(len(paragraph.runs) - 1, -1, -1):
            # Reverse loop so changes to run lengths don't affect the next
            # iteration.
            run_start = sum(run_lengths[:run_index])
            run_end = run_start + run_lengths[run_index]
            run_obj = paragraph.runs[run_index]
            for match_index in match_indices:
                if match_index < run_start or match_index >= run_end:
                    continue

                overflow = match_index + len(find) - run_end
                run_obj.text = run_obj.text[: match_index - run_start] + replace

                current_index = run_index + 1
                while overflow > 0:
                    next_run = paragraph.runs[current_index]
                    new_overflow = overflow - len(next_run.text)
                    next_run.text = next_run.text[overflow:]
                    overflow = new_overflow
                    current_index += 1

                break


def format_paragraph(  # noqa: PLR0913
    paragraph: docx_paragraph.Paragraph,
    *,
    line_spacing: float | None = None,
    spacing_before: float | None = None,
    spacing_after: float | None = None,
    bold: bool | None = None,
    italics: bool | None = None,
    font_size: int | None = None,
    font_rgb: tuple[int, int, int] | None = None,
    alignment: text.WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    """Formats a paragraph in a Word document.

    Args:
        paragraph: The paragraph to format.
        line_spacing: The line spacing of the paragraph.
        spacing_before: The spacing before the paragraph.
        spacing_after: The spacing after the paragraph.
        bold: Whether to bold the paragraph.
        italics: Whether to italicize the paragraph.
        italics: Whether to italicize the paragraph.
        font_size: The font size of the paragraph.
        font_rgb: The font color of the paragraph.
        alignment: The alignment of the paragraph.
    """
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing

    if alignment is not None:
        paragraph.alignment = alignment

    if spacing_before is not None:
        paragraph.paragraph_format.space_before = spacing_before

    if spacing_after is not None:
        paragraph.paragraph_format.space_after = spacing_after

    for run in paragraph.runs:
        if bold is not None:
            run.bold = bold
        if italics is not None:
            run.italic = italics
        if font_size is not None:
            run.font.size = font_size
        if font_rgb is not None:
            run.font.color.rgb = docx.shared.RGBColor(*font_rgb)


def format_cell(  # noqa: PLR0913, D417
    cell: table._Cell,
    *,
    line_spacing: float | None = None,
    spacing_before: float | None = None,
    spacing_after: float | None = None,
    bold: bool | None = None,
    italics: bool | None = None,
    font_size: int | None = None,
    font_rgb: tuple[int, int, int] | None = None,
    background_rgb: tuple[int, int, int] | None = None,
    alignment: text.WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    """Formats a cell in a Word table.

    Args:
        cell: The cell to format.
        line_spacing: The line spacing of the cell.
        spacing_before: The spacing before the cell.
        spacing_after: The spacing after the cell.
        bold: Whether to bold the cell.
        italics: Whether to italicize the cell.
        font_size: The font size of the cell.
        font_rgb: The font color of the cell.
        background_rgb: The background color of the cell.
    """
    for para in cell.paragraphs:
        format_paragraph(
            para,
            line_spacing=line_spacing,
            bold=bold,
            italics=italics,
            font_size=font_size,
            font_rgb=font_rgb,
            alignment=alignment,
            spacing_after=spacing_after,
            spacing_before=spacing_before,
        )

    if background_rgb is not None:
        shading = oxml.parse_xml(
            (r'<w:shd {} w:fill="' + f"{rgb_to_hex(*background_rgb)}" + r'"/>').format(
                ns.nsdecls("w"),
            ),
        )

        cell._tc.get_or_add_tcPr().append(shading)  # noqa: SLF001


def rgb_to_hex(r: int, g: int, b: int) -> str:
    """Converts RGB values to a hexadecimal color code.

    Args:
        r: The red component of the RGB color.
        g: The green component of the RGB color.
        b: The blue component of the RGB color.

    Returns:
        The hexadecimal color code representing the RGB color.
    """
    return f"#{r:02x}{g:02x}{b:02x}".upper()


def join_with_oxford_comma(
    items: list[str],
) -> str:
    """Joins a list of items with an Oxford comma.

    Args:
        items: The items to be joined.

    Returns:
        str: The joined string.
    """
    if len(items) == 0:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:  # noqa: PLR2004
        return f"{items[0]} and {items[1]}"
    return f"{', '.join(items[:-1])}, and {items[-1]}"


def ordinal_suffix(number: int | str) -> str:
    """Converts a number to its ordinal suffix.

    Args:
        number: The number to convert.

    Returns:
        str: The ordinal suffix of the number.
    """
    number = int(number)

    last_two_digits = number % 100
    if 11 <= last_two_digits <= 13:  # noqa: PLR2004
        return "th"

    last_digit = number % 10
    if last_digit == 1:
        return "st"
    if last_digit == 2:  # noqa: PLR2004
        return "nd"
    if last_digit == 3:  # noqa: PLR2004
        return "rd"
    return "th"


def read_subject_row(
    csv_file: fastapi.UploadFile,
    redcap_survey_identifier: int,
) -> pl.DataFrame:
    """Reads the subject row from the intake CSV file.

    All variables are interpreted as strings unless explicitly specified otherwise as
    the REDCap .csv is too inconsistent in its typing.

    Args:
        csv_file: The intake CSV file.
        redcap_survey_identifier: The REDCap survey identifier for the intake form.

    Returns:
        The subject row.

    Raises:
        HTTPException: If the subject is not found.
    """
    dtypes = {
        "age": pl.Float32,
        "birth_location": pl.Int8,
        "child_language1_fluency": pl.Int8,
        "child_language2_fluency": pl.Int8,
        "child_language3_fluency": pl.Int8,
        "childgender": pl.Int8,
        "dominant_hand": pl.Int8,
        "guardian_maritalstatus": pl.Int8,
        "guardian_relationship___1": pl.Int8,
        "iep": pl.Int8,
        "infanttemp_adapt": pl.Int8,
        "infanttemp1": pl.Int8,
        "language_spoken": pl.Int8,
        "opt_delivery": pl.Int8,
        "residing_number": pl.Int8,
        "pronouns": pl.Int8,
        "schooltype": pl.Int8,
    }

    faulty_people_in_home = 2
    for index in range(1, 11):
        dtypes[f"peopleinhome{index}_relation"] = pl.Int8
        if index != faulty_people_in_home:
            dtypes[f"peopleinhome{index}_relationship"] = pl.Int8
        else:
            dtypes["peopleinhome_relationship"] = pl.Int8

    intake_df = pl.read_csv(
        csv_file.file,
        infer_schema_length=0,
        dtypes=dtypes,
    )

    subject_df = intake_df.filter(
        intake_df["redcap_survey_identifier"] == redcap_survey_identifier,
    )
    if subject_df.height == 1:
        return subject_df

    if subject_df.height > 1:
        raise fastapi.HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Multiple patients found.",
        )

    raise fastapi.HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Patient not found.",
    )
